import docx
import os
import sys

def inspect_docx(file_path):
    sys.stdout.reconfigure(encoding='utf-8')
    print(f"Inspecting: {file_path}")
    if not os.path.exists(file_path):
        print("File does not exist!")
        return
    
    doc = docx.Document(file_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # Print first 20 paragraphs to see what it is
    print("\n--- FIRST 20 PARAGRAPHS ---")
    count = 0
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"{count}: {p.text[:100]} (Style: {p.style.name})")
            count += 1
            if count >= 30:
                break
                
    # List all headings
    print("\n--- HEADINGS ---")
    headings = []
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            headings.append((p.style.name, p.text))
            
    print(f"Total headings found: {len(headings)}")
    for style, text in headings:
        print(f"  {style}: {text}")

if __name__ == "__main__":
    inspect_docx(r"c:\Hub\DoAn\DoAn.docx")
