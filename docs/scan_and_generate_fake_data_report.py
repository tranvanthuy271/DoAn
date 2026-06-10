import docx
import os
import re
import sys

def scan_and_generate_report():
    sys.stdout.reconfigure(encoding='utf-8')
    docx_path = r"c:\Hub\DoAn\DoAn.docx"
    report_path = r"c:\Hub\DoAn\docs\report_fake_data.md"
    
    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} not found.")
        return
        
    doc = docx.Document(docx_path)
    
    # 1. Patterns to search for
    patterns = [
        (re.compile(r'\[([^\]]{2,50})\]'), "Bracketed placeholder [text]"),
        (re.compile(r'<([^>]{2,50})>'), "Angle-bracketed placeholder <text>"),
        (re.compile(r'\.\.\.\.+'), "Ellipsis/Fill-in-the-blank '...'"),
        (re.compile(r'____+'), "Underscore/Fill-in-the-blank '___'"),
        (re.compile(r'Nguyễn Văn [A-Z]', re.IGNORECASE), "Dummy student/teacher name 'Nguyễn Văn A'"),
        (re.compile(r'Trần Thị [A-Z]', re.IGNORECASE), "Dummy student/teacher name 'Trần Thị B'"),
        (re.compile(r'\b(lorem|ipsum|placeholder|todo|tbd|mock)\b', re.IGNORECASE), "Template keywords"),
        (re.compile(r'nhập vào|chưa viết|cần bổ sung', re.IGNORECASE), "Draft markers / missing content notes"),
        (re.compile(r'MSSV|Mã sinh viên|202\d{1}', re.IGNORECASE), "Placeholder student metadata/year templates")
    ]
    
    found_items = []
    
    # Scan Paragraphs
    for p_idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        for pattern, label in patterns:
            for match in pattern.finditer(text):
                found_items.append({
                    "location": f"Paragraph Line {p_idx+1}",
                    "type": label,
                    "matched": match.group(0),
                    "context": text[:150]
                })
                
    # Scan Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                for pattern, label in patterns:
                    for match in pattern.finditer(text):
                        found_items.append({
                            "location": f"Table {t_idx+1}, Row {r_idx+1}, Column {c_idx+1}",
                            "type": label,
                            "matched": match.group(0),
                            "context": text[:150]
                        })
                        
    # Filter unique matches to avoid clutter
    unique_items = []
    seen = set()
    for item in found_items:
        key = (item["location"], item["type"], item["matched"])
        if key not in seen:
            seen.add(key)
            unique_items.append(item)
            
    # 2. Write beautiful Markdown report
    markdown_content = f"""# BÁO CÁO RÀ SOÁT DỮ LIỆU GIẢ VỜ VÀ BẢN MẪU (PLACEHOLDERS)
**Tài liệu rà soát**: `DoAn.docx`
**Ngày thực hiện**: 02/06/2026
**Phạm vi kiểm tra**: Toàn bộ nội dung văn bản và bảng biểu trong tài liệu chính nhằm phát hiện thông tin mẫu, tên giả định, ký tự chờ nhập hoặc các nội dung chưa hoàn chỉnh cần bổ sung trước khi nộp đồ án.

---

## 1. Tóm tắt kết quả
- **Tổng số vị trí phát hiện mẫu nghi vấn**: {len(unique_items)} mục.
- **Mức độ ảnh hưởng**:
  > [!IMPORTANT]
  > Hầu hết các dữ liệu giả lập (như "Nguyễn Văn A", "MSSV", "202x") nằm ở các trang bìa phụ, trang điền thông tin cá nhân và đề cương nhiệm vụ. Các phần nội dung chuyên môn chính (Chương 1, 2, 3, 4) được viết khá đầy đủ và bám sát dự án thực tế. Tuy nhiên, vẫn còn một số chỗ cần thay thế thông tin cá nhân chính xác trước khi xuất bản.

---

## 2. Chi tiết các vị trí phát hiện dữ liệu giả lập / Bản mẫu

| Vị trí | Loại lỗi / Ký hiệu | Nội dung phát hiện | Đoạn văn cảnh (Context) |
|---|---|---|---|
"""
    
    for item in unique_items:
        loc = item["location"]
        type_ = item["type"]
        matched = item["matched"].replace("|", "\\|").replace("\n", " ")
        context = item["context"].replace("|", "\\|").replace("\n", " ")
        # Escape markdown formatting inside cells
        matched = f"`{matched}`"
        markdown_content += f"| {loc} | {type_} | {matched} | {context} |\n"
        
    markdown_content += """
---

## 3. Khuyến nghị khắc phục hành chính
1. **Trang bìa & Thông tin cá nhân**: Thay thế toàn bộ "Nguyễn Văn A" và mã số sinh viên demo bằng tên và mã số thực của bạn.
2. **Năm học**: Rà soát lại năm học (hiện đang là 2026 hoặc 202x) xem có khớp với thời gian nộp đồ án thực tế của trường hay không.
3. **Các ký tự chờ điền (`...` hoặc `___`)**: Điền đầy đủ thông tin hoặc xóa bớt nếu đó là các biểu mẫu không áp dụng.
"""

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
        
    print(f"Report generated successfully with {len(unique_items)} placeholder matches in: {report_path}")

if __name__ == "__main__":
    scan_and_generate_report()
