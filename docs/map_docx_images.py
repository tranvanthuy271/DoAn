import os
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r"c:\Hub\DoAn\DoAn.docx"

def map_images():
    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} does not exist.")
        return

    doc = Document(docx_path)
    print(f"Loaded {docx_path}. Mapping images to captions...")

    # We want to find all relationships to map rId -> target file
    part = doc.part
    rel_map = {}
    for rId, rel in part.rels.items():
        if "image" in rel.target_ref:
            rel_map[rId] = os.path.basename(rel.target_ref)

    print(f"Found {len(rel_map)} image relationships in document parts.")

    # Iterate through all paragraphs and table cells to find inline drawings
    drawing_count = 0
    
    # Helper function to check if a paragraph has drawings and return their filenames
    def get_para_images(p):
        filenames = []
        # Find all drawing elements in this paragraph's XML
        # Drawings are represented by <w:drawing> elements
        drawings = p._element.xpath('.//w:drawing')
        for drawing in drawings:
            # Inside drawing, find embed attribute in blip element <a:blip r:embed="rId...">
            blips = drawing.xpath('.//a:blip')
            for blip in blips:
                embed_id = blip.get(qn('r:embed'))
                if embed_id and embed_id in rel_map:
                    filenames.append(rel_map[embed_id])
        return filenames

    # We will scan paragraphs in order and keep track of recent paragraph text to find captions
    paragraph_log = []
    
    print("\n--- Image Mapping ---")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        paragraph_log.append((idx, text))
        
        # Check if paragraph has images
        img_files = get_para_images(p)
        if img_files:
            drawing_count += len(img_files)
            # Find the nearest caption before or after this paragraph
            nearest_caption = "No caption found"
            
            # Search backwards for a caption (up to 10 paragraphs back)
            for j in range(len(paragraph_log) - 1, max(-1, len(paragraph_log) - 12), -1):
                t = paragraph_log[j][1]
                if t.startswith("Hình") or t.startswith("Hình") or t.startswith("Sơ đồ") or t.startswith("Bảng") or "hình 2." in t.lower() or "hình 2." in t.lower():
                    nearest_caption = f"Previous caption: '{t}' (para {paragraph_log[j][0]+1})"
                    break
            
            # If no previous caption, look ahead (up to 10 paragraphs ahead)
            if nearest_caption == "No caption found":
                for k in range(idx + 1, min(len(doc.paragraphs), idx + 12)):
                    t = doc.paragraphs[k].text.strip()
                    if t.startswith("Hình") or t.startswith("Hình") or t.startswith("Sơ đồ") or t.startswith("Bảng") or "hình 2." in t.lower() or "hình 2." in t.lower():
                        nearest_caption = f"Next caption: '{t}' (para {k+1})"
                        break
            
            # Print the match
            print(f"Para {idx+1}: {img_files} -> {nearest_caption}")
            if text:
                print(f"   [Paragraph text]: '{text}'")

    print(f"\nTotal drawings identified in paragraphs: {drawing_count}")

if __name__ == "__main__":
    map_images()
