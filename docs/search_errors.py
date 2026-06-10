import docx
import sys
import os

def search_errors(file_path):
    sys.stdout.reconfigure(encoding='utf-8')
    if not os.path.exists(file_path):
        print(f"File {file_path} not found!")
        return
    
    doc = docx.Document(file_path)
    
    search_terms = {
        "IEnemyService": "IEnemyService (Error 1)",
        "_userRepository.FindByUsernameAsync": "_userRepository.FindByUsernameAsync (Error 2)",
        "ZoneApiKey": "ZoneApiKey (Error 3)",
        "[Authorize(Roles = \"GameServer\")]": "[Authorize(Roles = \"GameServer\")] (Error 4)",
        "ErrorHandlingMiddleware": "ErrorHandlingMiddleware (Error 5/13)",
        "FixedTimeEquals": "FixedTimeEquals (Error 3)",
        "3306 (chỉ mạng nội bộ": "3306 (chỉ mạng nội bộ) (Error 6)",
        "7777 UDP": "7777 UDP (Error 7)",
        "ConnectionStrings__DefaultConnection": "ConnectionStrings__DefaultConnection (Error 8)",
        "git pull origin main": "git pull origin main (Error 9)",
        "ConnectionApprovalResponse.Reason": "ConnectionApprovalResponse.Reason (Error 14)"
    }
    
    print("Searching for errors in paragraphs...")
    found_count = 0
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        for term, desc in search_terms.items():
            if term in text:
                print(f"[FOUND Paragraph] Line {idx+1} | {desc}\n  Content: {text[:150]}\n")
                found_count += 1
                
    print("Searching for errors in tables...")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                for term, desc in search_terms.items():
                    if term in text:
                        print(f"[FOUND Table {t_idx+1}, R{r_idx+1}, C{c_idx+1}] | {desc}\n  Content: {text[:150]}\n")
                        found_count += 1
                        
    print(f"Search complete. Total error references found: {found_count}")

if __name__ == "__main__":
    search_errors(r"c:\Hub\DoAn\DoAn.docx")
