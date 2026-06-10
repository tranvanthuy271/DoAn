import docx
import sys
import os

def search_classes(file_path):
    sys.stdout.reconfigure(encoding='utf-8')
    if not os.path.exists(file_path):
        print(f"File {file_path} not found!")
        return
    
    doc = docx.Document(file_path)
    
    classes = [
        "EnemyController", "AuthController", "ZoneApiKeyAuthenticationHandler", 
        "ErrorHandlingMiddleware", "deploy.sh", "ConnectionApprovalResponse",
        "SecureEquals", "IEnemyService", "_userRepository", "BCrypt",
        "FixedTimeEquals", "ZoneApiKey"
    ]
    
    print("Searching classes in paragraphs...")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        for cls in classes:
            if cls in text:
                print(f"[Paragraph] Line {idx+1} | Keyword: {cls}\n  Content: {text[:150]}\n")
                
    print("Searching classes in tables...")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                for cls in classes:
                    if cls in text:
                        print(f"[Table {t_idx+1}, R{r_idx+1}, C{c_idx+1}] | Keyword: {cls}\n  Content: {text[:150]}\n")

if __name__ == "__main__":
    search_classes(r"c:\Hub\DoAn\DoAn.docx")
