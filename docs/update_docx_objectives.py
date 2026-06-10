import docx
import os
import sys

def update_docx(file_path):
    print(f"Loading document: {file_path}")
    if not os.path.exists(file_path):
        print(f"File {file_path} not found!")
        return False
        
    doc = docx.Document(file_path)
    modified = False
    
    # 1. Search in paragraphs
    for idx, p in enumerate(doc.paragraphs):
        # Target paragraph text: "Mục tiêu tổng quát của đề tài là xây dựng một nguyên mẫu trò chơi Mutants Arena có thể chơi được, có đầy đủ các thành phần cốt lõi của một game 2D Action RPG và..."
        target_str = "xây dựng một nguyên mẫu trò chơi Mutants Arena có thể chơi được, có đầy đủ các thành phần cốt lõi của một game 2D Action RPG"
        if target_str in p.text:
            print(f"Found match in paragraph {idx}: {p.text[:100]}...")
            # Let's inspect runs first
            print("Runs inside paragraph:")
            for run_idx, run in enumerate(p.runs):
                print(f"  Run {run_idx}: '{run.text}'")
            
            # Replace the text. To preserve formatting if there are multiple runs, we can replace the text within the runs,
            # but since it's standard Normal text, we can also reconstruct or replace carefully.
            # Let's perform a simple replace on the entire paragraph text. Since the style is Normal,
            # this is safe.
            old_text = p.text
            new_text = old_text.replace(
                "xây dựng một nguyên mẫu trò chơi Mutants Arena có thể chơi được, có đầy đủ các thành phần cốt lõi của một game 2D Action RPG",
                "xây dựng một nguyên mẫu trò chơi multiplayer Mutants Arena có thể chơi được, có đầy đủ các thành phần cốt lõi của một game 2D Action RPG"
            )
            
            # Let's clear the paragraph and set the text
            p.text = new_text
            print(f"Updated paragraph {idx} text to: {p.text[:120]}...")
            modified = True
            
    # 2. Also search in tables just in case
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    target_str = "xây dựng một nguyên mẫu trò chơi Mutants Arena có thể chơi được, có đầy đủ các thành phần cốt lõi của một game 2D Action RPG"
                    if target_str in p.text:
                        print(f"Found match in table {t_idx}, row {r_idx}, cell {c_idx}, paragraph {p_idx}")
                        p.text = p.text.replace(
                            "xây dựng một nguyên mẫu trò chơi Mutants Arena có thể chơi được, có đầy đủ các thành phần cốt lõi của một game 2D Action RPG",
                            "xây dựng một nguyên mẫu trò chơi multiplayer Mutants Arena có thể chơi được, có đầy đủ các thành phần cốt lõi của một game 2D Action RPG"
                        )
                        modified = True

    if modified:
        doc.save(file_path)
        print(f"Successfully saved updated document to {file_path}\n")
        return True
    else:
        print(f"No match found in {file_path}\n")
        return False

if __name__ == "__main__":
    sys.stdout.reconfigure(encoding='utf-8')
    res1 = update_docx(r"c:\Hub\DoAn\DoAn.docx")
    res2 = update_docx(r"c:\Hub\DoAn\docs\DoAn.docx")
    
    if res1 or res2:
        print("Objective update completed!")
    else:
        print("No document was updated.")
