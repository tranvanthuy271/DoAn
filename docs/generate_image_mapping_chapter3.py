import os
import sys
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r"c:\Hub\DoAn\DoAn.docx"
doc = Document(docx_path)
rel_map = {rId: os.path.basename(rel.target_ref) for rId, rel in doc.part.rels.items() if 'image' in rel.target_ref}

output_path = r"c:\Hub\DoAn\docs\image_mapping_chapter3.md"
with open(output_path, 'w', encoding='utf-8') as f:
    f.write("# Chapter 3 Image Mapping from DoAn.docx\n\n")
    for i in range(970, 1200):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            text = p.text.strip()
            img_files = []
            for d in p._element.xpath('.//w:drawing'):
                for b in d.xpath('.//a:blip'):
                    embed_id = b.get(qn('r:embed'))
                    if embed_id and embed_id in rel_map:
                        img_files.append(rel_map[embed_id])
            if text or img_files:
                f.write(f"**Para {i+1}**: `{img_files}` | Text: {text}\n\n")

print("Generated docs/image_mapping_chapter3.md successfully!")
