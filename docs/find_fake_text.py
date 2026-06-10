import docx
import sys
import re
import os

def find_fake_text(file_path):
    sys.stdout.reconfigure(encoding='utf-8')
    if not os.path.exists(file_path):
        print(f"File {file_path} not found!")
        return
    
    doc = docx.Document(file_path)
    
    patterns = [
        r"Nguyễn Văn [A-Z]",
        r"Trần Thị [A-Z]",
        r"MSSV",
        r"Mã sinh viên",
        r"\b\d{6,10}\b", # potential dummy student IDs
        r"202\d", # potential template years like 2026, 2024, etc.
        r"Giảng viên hướng dẫn",
        r"Sinh viên thực hiện",
        r"Lớp:",
        r"Khoa:",
        r"Trường Đại học",
        r"Đề tài:"
    ]
    
    print("Searching for potential template/personal placeholders...")
    found = []
    
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        
        for pat in patterns:
            if re.search(pat, text, re.IGNORECASE):
                found.append((f"Paragraph Line {idx+1}", pat, text[:150]))
                
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                for pat in patterns:
                    if re.search(pat, text, re.IGNORECASE):
                        found.append((f"Table {t_idx+1}, R{r_idx+1}, C{c_idx+1}", pat, text[:150]))
                        
    # Remove duplicates
    unique_found = []
    seen = set()
    for item in found:
        key = (item[0], item[1])
        if key not in seen:
            seen.add(key)
            unique_found.append(item)
            
    print(f"Total potential template details found: {len(unique_found)}")
    for loc, pat, text in unique_found[:100]:
        print(f"[{loc}] | Pattern: {pat}\n  Content: {text}\n")

if __name__ == "__main__":
    find_fake_text(r"c:\Hub\DoAn\DoAn.docx")
