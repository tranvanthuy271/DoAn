import docx
import sys

def print_paragraphs(file_path):
    sys.stdout.reconfigure(encoding='utf-8')
    doc = docx.Document(file_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    for idx in range(935, min(975, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        print(f"Line {idx}: [{p.style.name}] {p.text}")

if __name__ == "__main__":
    print_paragraphs(r"c:\Hub\DoAn\DoAn.docx")
