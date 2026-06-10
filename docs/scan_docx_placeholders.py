import docx
import os
import re
import sys

def scan_placeholders(file_path):
    sys.stdout.reconfigure(encoding='utf-8')
    if not os.path.exists(file_path):
        print(f"File {file_path} not found!")
        return
    
    doc = docx.Document(file_path)
    
    print(f"Scanning document: {file_path}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    
    placeholders = []
    
    # Common placeholder regex patterns
    # 1. Square brackets [something]
    bracket_pattern = re.compile(r'\[([^\]]{2,50})\]')
    # 2. Angle brackets <something>
    angle_pattern = re.compile(r'<([^>]{2,50})>')
    # 3. Triple dots or underscores for fill-in-the-blanks
    dots_pattern = re.compile(r'\.\.\.\.+')
    underscore_pattern = re.compile(r'____+')
    # 4. Keyword matches (case-insensitive)
    keywords = ['lorem', 'ipsum', 'placeholder', 'todo', 'tbd', 'fake', 'mock', 'nhập vào', 'chưa viết', 'cần bổ sung']
    
    # Scan Paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        # Match bracket
        for m in bracket_pattern.finditer(text):
            placeholders.append(('paragraph', idx, f"Line {idx+1}", f"Bracket match: '{m.group(0)}'", text[:150]))
            
        # Match angle
        for m in angle_pattern.finditer(text):
            placeholders.append(('paragraph', idx, f"Line {idx+1}", f"Angle match: '{m.group(0)}'", text[:150]))
            
        # Match dots
        if dots_pattern.search(text):
            placeholders.append(('paragraph', idx, f"Line {idx+1}", "Multiple dots '...'", text[:150]))
            
        # Match underscores
        if underscore_pattern.search(text):
            placeholders.append(('paragraph', idx, f"Line {idx+1}", "Multiple underscores '___'", text[:150]))
            
        # Match keywords
        lower_text = text.lower()
        for kw in keywords:
            if kw in lower_text:
                placeholders.append(('paragraph', idx, f"Line {idx+1}", f"Keyword '{kw}' match", text[:150]))
                
    # Scan Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                
                # Match bracket
                for m in bracket_pattern.finditer(text):
                    placeholders.append(('table', t_idx, f"Table {t_idx+1}, Row {r_idx+1}, Col {c_idx+1}", f"Bracket match: '{m.group(0)}'", text[:150]))
                
                # Match angle
                for m in angle_pattern.finditer(text):
                    placeholders.append(('table', t_idx, f"Table {t_idx+1}, Row {r_idx+1}, Col {c_idx+1}", f"Angle match: '{m.group(0)}'", text[:150]))
                
                # Match dots
                if dots_pattern.search(text):
                    placeholders.append(('table', t_idx, f"Table {t_idx+1}, Row {r_idx+1}, Col {c_idx+1}", "Multiple dots '...'", text[:150]))
                
                # Match underscores
                if underscore_pattern.search(text):
                    placeholders.append(('table', t_idx, f"Table {t_idx+1}, Row {r_idx+1}, Col {c_idx+1}", "Multiple underscores '___'", text[:150]))
                
                # Match keywords
                lower_text = text.lower()
                for kw in keywords:
                    if kw in lower_text:
                        placeholders.append(('table', t_idx, f"Table {t_idx+1}, Row {r_idx+1}, Col {c_idx+1}", f"Keyword '{kw}' match", text[:150]))
                        
    # Remove duplicates
    unique_placeholders = []
    seen = set()
    for item in placeholders:
        key = (item[0], item[1], item[2], item[3])
        if key not in seen:
            seen.add(key)
            unique_placeholders.append(item)
            
    print(f"\nFound {len(unique_placeholders)} potential placeholder items:")
    for type_, loc_id, loc_str, reason, snippet in unique_placeholders:
        print(f"[{type_.upper()}] {loc_str} | Reason: {reason}\n  Content: {snippet}\n")
        
    return unique_placeholders

if __name__ == "__main__":
    scan_placeholders(r"c:\Hub\DoAn\DoAn.docx")
