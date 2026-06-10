import os
import sys
from docx import Document
from docx.oxml.ns import qn

def inspect():
    sys.stdout.reconfigure(encoding='utf-8')
    docx_path = r"c:\Hub\DoAn\DoAn.docx"
    doc = Document(docx_path)
    
    rel_map = {}
    for rId, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            rel_map[rId] = os.path.basename(rel.target_ref)
            
    print(f"Inspecting paragraphs index 720 to 745:")
    for i in range(720, 745):
        if i >= len(doc.paragraphs):
            break
        p = doc.paragraphs[i]
        text = p.text.strip()
        
        imgs = []
        blips = p._element.xpath('.//a:blip')
        for blip in blips:
            embed_id = blip.get(qn('r:embed'))
            if embed_id and embed_id in rel_map:
                imgs.append(rel_map[embed_id])
                
        if text or imgs:
            print(f"Index {i} (Para {i+1}): text='{text}' | images={imgs}")

if __name__ == "__main__":
    inspect()
