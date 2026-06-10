import docx
import sys
import os

def inspect_csharp(file_path):
    sys.stdout.reconfigure(encoding='utf-8')
    if not os.path.exists(file_path):
        print(f"File {file_path} not found!")
        return
    
    doc = docx.Document(file_path)
    
    # We will search for occurrences of text that resembles the code snippets of:
    # 1. EnemyController (GetAllEnemies)
    # 2. AuthController (Login)
    # 3. ZoneApiKeyAuthenticationHandler (SecureEquals / FixedTimeEquals)
    # 4. ErrorHandlingMiddleware
    
    print("--- IN-DEPTH CODE INSPECTION ---")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        if "GetAllEnemies" in text or "GetEnemiesByLevel" in text:
            print(f"[Line {idx}] {text[:200]}")
            # print surrounding paragraphs
            for offset in range(-2, 12):
                if 0 <= idx + offset < len(doc.paragraphs):
                    print(f"  [{idx+offset}]: {doc.paragraphs[idx+offset].text.strip()[:150]}")
            print("-" * 50)
            
        if "Login" in text and "ActionResult" in text:
            print(f"[Line {idx}] {text[:200]}")
            for offset in range(-2, 12):
                if 0 <= idx + offset < len(doc.paragraphs):
                    print(f"  [{idx+offset}]: {doc.paragraphs[idx+offset].text.strip()[:150]}")
            print("-" * 50)

if __name__ == "__main__":
    inspect_csharp(r"c:\Hub\DoAn\DoAn.docx")
