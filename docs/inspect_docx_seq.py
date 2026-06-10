import os
import sys
from docx import Document
from docx.oxml.ns import qn

def inspect():
    sys.stdout.reconfigure(encoding='utf-8')
    docx_path = r"c:\Hub\DoAn\DoAn.docx"
    doc = Document(docx_path)
    
    # Get relations
    rel_map = {}
    for rId, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            rel_map[rId] = os.path.basename(rel.target_ref)
            
    print(f"Inspecting paragraphs 750 to 795 of {docx_path}:")
    for i in range(750, 795):
        if i >= len(doc.paragraphs):
            break
        p = doc.paragraphs[i]
        text = p.text.strip()
        
        # Check drawings
        imgs = []
        blips = p._element.xpath('.//a:blip')
        for blip in blips:
            embed_id = blip.get(qn('r:embed'))
            if embed_id and embed_id in rel_map:
                imgs.append(rel_map[embed_id])
                
        if text or imgs:
            print(f"Para {i+1}: text='{text}' | images={imgs}")

if __name__ == "__main__":
    inspect()
