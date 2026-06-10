import docx
import sys
import os

def extract_text(file_path, output_path):
    sys.stdout.reconfigure(encoding='utf-8')
    if not os.path.exists(file_path):
        print(f"File {file_path} not found!")
        return
        
    doc = docx.Document(file_path)
    print(f"Extracting {file_path}...")
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"=== DOCUMENT EXTRACTED: {file_path} ===\n")
        f.write(f"Total Paragraphs: {len(doc.paragraphs)}\n")
        f.write(f"Total Tables: {len(doc.tables)}\n\n")
        
        f.write("=== PARAGRAPHS ===\n")
        for idx, p in enumerate(doc.paragraphs):
            style = p.style.name if p.style else "NoStyle"
            f.write(f"[P_{idx}] ({style}) {p.text}\n")
            
        f.write("\n=== TABLES ===\n")
        for t_idx, table in enumerate(doc.tables):
            f.write(f"\n--- Table {t_idx} ---\n")
            for r_idx, row in enumerate(table.rows):
                row_cells = []
                for c_idx, cell in enumerate(row.cells):
                    row_cells.append(cell.text.strip().replace("\n", " | "))
                f.write(f"[T_{t_idx}_R{r_idx}] " + " || ".join(row_cells) + "\n")
                
    print(f"Extraction complete. Saved to {output_path}")

if __name__ == "__main__":
    extract_text(r"c:\Hub\DoAn\DoAn.docx", r"c:\Hub\DoAn\docs\DoAn_extracted_text.txt")
