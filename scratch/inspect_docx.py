import docx

doc = docx.Document("DoAn.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")

# Print the first 50 headings
with open("scratch/docx_headings.txt", "w", encoding="utf-8") as f:
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") or p.text.isupper() or len(p.text) < 100 and any(keyword in p.text.lower() for keyword in ["chương", "mục lục", "tổng quan", "phân tích", "thiết kế", "xây dựng", "thực nghiệm", "kết luận"]):
            f.write(f"Para {i} | Style: {p.style.name} | Text: {p.text}\n")

print("Docx headings extraction completed.")
